from docx import Document


def preview_section_text(doc_path: str, section: str, table_index: int | None = None) -> str:
    doc = Document(doc_path)

    # ---------- helpers (same logic as ExperienceEditor) ----------
    def body_children():
        return list(doc._body._element)

    def paragraph_from_elem(elem):
        for p in doc.paragraphs:
            if p._p is elem:
                return p
        return None

    def find_table_pos(ti: int) -> int:
        tbl_elem = doc.tables[ti]._tbl
        children = body_children()
        for i, child in enumerate(children):
            if child is tbl_elem:
                return i
        raise ValueError("Table not found in document body.")

    def is_bullet_paragraph(p) -> bool:
        # true numbering/bullets in XML
        if p._p is not None and p._p.pPr is not None and p._p.pPr.numPr is not None:
            return True

        # style heuristics
        try:
            style_name = (p.style.name or "").lower()
        except Exception:
            style_name = ""
        if "list" in style_name or "bullet" in style_name:
            return True

        # literal bullet char
        return (p.text or "").lstrip().startswith("•")

    def bullet_texts_after_table(ti: int) -> list[str]:
        children = body_children()
        pos = find_table_pos(ti)

        bullets = []
        started = False

        for child in children[pos + 1:]:
            if child.tag.endswith("}tbl"):
                break
            if not child.tag.endswith("}p"):
                continue

            p_obj = paragraph_from_elem(child)
            if p_obj is None:
                continue

            txt = (p_obj.text or "").strip()
            if not txt:
                continue

            if is_bullet_paragraph(p_obj):
                bullets.append(txt)
                started = True
            else:
                if started:
                    break

        return bullets
    # -------------------------------------------------------------

    # ✅ Per-entry preview for EXPERIENCE/PROJECTS
    if section in ("EXPERIENCE", "PROJECTS") and table_index is not None:
        if table_index < 0 or table_index >= len(doc.tables):
            return f"(Invalid table_index: {table_index}. Download to verify.)"

        tbl = doc.tables[table_index]
        row = tbl.rows[0]

        left = row.cells[0].text.strip() if len(row.cells) > 0 else ""
        right = row.cells[1].text.strip() if len(row.cells) > 1 else ""

        bullets = bullet_texts_after_table(table_index)

        out = []
        if left or right:
            out.append(f"{left} | {right}".strip(" |"))

        if bullets:
            out.extend(bullets)
        else:
            out.append("(No bullets found under this entry. Download to verify.)")

        return "\n".join(out)

    # ✅ generic fallback for other sections (keep your existing logic)
    lines = []
    hit = False
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        if txt == section:
            hit = True
            continue
        if hit:
            if txt.isupper() and len(txt) < 40:
                break
            lines.append(txt)
            if len(lines) >= 10:
                break

    if not lines:
        return f"(No preview text found for {section}. Download to verify.)"
    return "\n".join(lines)