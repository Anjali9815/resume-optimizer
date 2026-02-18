from docx import Document
import re

SECTION_REGEX = re.compile(r"^[A-Z][A-Z\s&]{2,}$")  # SUMMARY, EDUCATION, TECHNICAL SKILLS, etc.


def detect_headers(doc_path: str) -> list[str]:
    doc = Document(doc_path)
    headers = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        # ignore name header (could also be caps); still ok if included
        if SECTION_REGEX.match(txt):
            if txt not in headers:
                headers.append(txt)
    return headers


def scan_tables_text_date(doc_path: str) -> list[int]:
    """
    Returns table indices that look like 2 columns: left=text, right=date.
    You already validated these in your scan script.
    """
    doc = Document(doc_path)
    good = []
    for ti, tbl in enumerate(doc.tables):
        if len(tbl.rows) < 1:
            continue
        row = tbl.rows[0]
        if len(row.cells) != 2:
            continue
        left = " ".join(p.text.strip() for p in row.cells[0].paragraphs if p.text.strip()).strip()
        right = " ".join(p.text.strip() for p in row.cells[1].paragraphs if p.text.strip()).strip()
        if left and right:
            good.append(ti)
    return good


def section_table_map(headers: list[str], table_indices: list[int]) -> dict[str, list[int]]:
    """
    Your current resume uses:
    EDUCATION table: 0
    EXPERIENCE tables: 1,2,3
    PROJECTS tables: 4,5,6

    For now we keep this deterministic mapping (fast + reliable).
    Later we can auto-detect boundaries by header positions.
    """
    m = {}

    # only add if those headers exist
    if "EDUCATION" in headers and 0 in table_indices:
        m["EDUCATION"] = [0]

    if "EXPERIENCE" in headers:
        m["EXPERIENCE"] = [i for i in table_indices if i in (1, 2, 3)]

    if "PROJECTS" in headers:
        m["PROJECTS"] = [i for i in table_indices if i in (4, 5, 6)]

    return m
