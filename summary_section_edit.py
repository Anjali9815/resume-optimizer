# summary_section_edit.py
from docx import Document
import re


class SummaryEditor:
    HEADING_RE = re.compile(r"^[A-Z0-9 &/\-]+$")

    def __init__(self, resume_path: str):
        self.resume_path = resume_path
        self.doc = Document(resume_path)

    def _is_caps_header(self, text: str) -> bool:
        t = (text or "").strip()
        if not t:
            return False
        letters = [c for c in t if c.isalpha()]
        return bool(letters) and t == t.upper() and bool(self.HEADING_RE.match(t))

    def _find_section_range(self, header: str):
        header = header.strip().upper()
        paras = self.doc.paragraphs

        heading_idx = None
        for i, p in enumerate(paras):
            if (p.text or "").strip().upper() == header:
                heading_idx = i
                break
        if heading_idx is None:
            raise ValueError(f"{header} not found.")

        start = heading_idx + 1
        end = start
        while end < len(paras):
            t = (paras[end].text or "").strip()
            if t and self._is_caps_header(t):
                break
            end += 1
        return start, end

    def get_current(self) -> str:
        start, end = self._find_section_range("SUMMARY")
        lines = []
        for p in self.doc.paragraphs[start:end]:
            if (p.text or "").strip():
                lines.append(p.text.strip())
        return "\n".join(lines).strip()

    def _clear_paragraph(self, paragraph):
        if paragraph.runs:
            for r in paragraph.runs:
                r.text = ""
        else:
            paragraph.text = ""

    def _set_text_keep_first_run_format(self, paragraph, text: str):
        if paragraph.runs:
            r0 = paragraph.runs[0]
            bold, italic, underline = r0.bold, r0.italic, r0.underline
            font_name, font_size = r0.font.name, r0.font.size

            for r in paragraph.runs:
                r.text = ""

            r0.text = text
            r0.bold, r0.italic, r0.underline = bold, italic, underline
            r0.font.name, r0.font.size = font_name, font_size
        else:
            paragraph.add_run(text)

    def update(self, new_summary: str):
        start, end = self._find_section_range("SUMMARY")
        block = self.doc.paragraphs[start:end]
        if not block:
            raise ValueError("No SUMMARY content found.")

        # Keep it in ONE paragraph to preserve section layout
        one_para = " ".join(new_summary.splitlines()).strip()
        if not one_para:
            raise ValueError("Summary cannot be empty.")

        self._set_text_keep_first_run_format(block[0], one_para)

        # Keep remaining paragraphs as empty spacers (no deletion!)
        for p in block[1:]:
            self._clear_paragraph(p)

    def save(self, output_path: str):
        self.doc.save(output_path)
