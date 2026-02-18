# skills_edit.py
from docx import Document
import re


class SkillsEditor:
    HEADING_RE = re.compile(r"^[A-Z0-9 &/\-]+$")

    def __init__(self, resume_path: str):
        self.resume_path = resume_path
        self.doc = Document(resume_path)

    def _is_caps_header(self, text: str) -> bool:
        t = (text or "").strip()
        if not t:
            return False
        letters = [c for c in t if c.isalpha()]
        return bool(letters) and t == t.upper() and bool(self.HEADING_RE.match(t))

    def _find_heading_idx(self, header: str) -> int:
        header = header.strip().upper()
        for i, p in enumerate(self.doc.paragraphs):
            if (p.text or "").strip().upper() == header:
                return i
        raise ValueError(f"Header '{header}' not found.")

    def _find_section_range(self, header: str) -> tuple[int, int]:
        heading_idx = self._find_heading_idx(header)
        start = heading_idx + 1
        end = start
        paras = self.doc.paragraphs
        while end < len(paras):
            t = (paras[end].text or "").strip()
            if t and self._is_caps_header(t):
                break
            end += 1
        return start, end

    def _delete_paragraph(self, paragraph) -> None:
        el = paragraph._element
        el.getparent().remove(el)
        paragraph._p = paragraph._element = None

    def _snapshot_template(self, paragraphs):
        """
        Snapshot paragraph + run formatting BEFORE deleting anything.
        Returns dict with style + paragraph_format attrs + run font attrs.
        """
        # pick first paragraph as base
        base_p = paragraphs[0]

        # find a run to copy font from (first run in block)
        base_run = None
        for p in paragraphs:
            if p.runs:
                base_run = p.runs[0]
                break

        pf = base_p.paragraph_format

        tpl = {
            "style": base_p.style,

            # paragraph format snapshot
            "alignment": pf.alignment,
            "left_indent": pf.left_indent,
            "right_indent": pf.right_indent,
            "first_line_indent": pf.first_line_indent,
            "space_before": pf.space_before,
            "space_after": pf.space_after,
            "line_spacing": pf.line_spacing,
            "keep_together": pf.keep_together,
            "keep_with_next": pf.keep_with_next,
            "page_break_before": pf.page_break_before,
            "widow_control": pf.widow_control,

            # run/font snapshot (optional)
            "run": None
        }

        if base_run is not None:
            tpl["run"] = {
                "bold": base_run.bold,
                "italic": base_run.italic,
                "underline": base_run.underline,
                "font_name": base_run.font.name,
                "font_size": base_run.font.size,
                "font_rgb": base_run.font.color.rgb if base_run.font.color else None,
            }

        return tpl

    def _apply_paragraph_format(self, p, tpl):
        pf = p.paragraph_format
        pf.alignment = tpl["alignment"]
        pf.left_indent = tpl["left_indent"]
        pf.right_indent = tpl["right_indent"]
        pf.first_line_indent = tpl["first_line_indent"]
        pf.space_before = tpl["space_before"]
        pf.space_after = tpl["space_after"]
        pf.line_spacing = tpl["line_spacing"]
        pf.keep_together = tpl["keep_together"]
        pf.keep_with_next = tpl["keep_with_next"]
        pf.page_break_before = tpl["page_break_before"]
        pf.widow_control = tpl["widow_control"]

    def _add_line_with_template(self, insert_before_p, text: str, tpl):
        """
        Create a new paragraph with the same style/formatting and insert it
        right before insert_before_p (a docx paragraph object).
        """
        p = self.doc.add_paragraph("", style=tpl["style"])
        insert_before_p._p.addprevious(p._p)

        self._apply_paragraph_format(p, tpl)

        run_tpl = tpl["run"]
        if run_tpl:
            r = p.add_run(text)
            r.bold = run_tpl["bold"]
            r.italic = run_tpl["italic"]
            r.underline = run_tpl["underline"]
            r.font.name = run_tpl["font_name"]
            r.font.size = run_tpl["font_size"]
            if run_tpl["font_rgb"]:
                r.font.color.rgb = run_tpl["font_rgb"]
        else:
            p.add_run(text)

    def get_current_lines(self) -> list[str]:
        start, end = self._find_section_range("TECHNICAL SKILLS")
        return [(p.text or "") for p in self.doc.paragraphs[start:end]]

    def replace_whole_section(self, pasted_text: str):
        start, end = self._find_section_range("TECHNICAL SKILLS")
        paras = self.doc.paragraphs

        if end >= len(paras):
            raise ValueError("Couldn't find next section after TECHNICAL SKILLS.")

        next_heading_para = paras[end]  # e.g., EXPERIENCE heading

        old_block = paras[start:end]
        if not old_block:
            raise ValueError("No TECHNICAL SKILLS content found.")

        # count trailing blank paragraphs inside old block (spacing gap)
        trailing_blank_count = 0
        for p in reversed(old_block):
            if (p.text or "").strip() == "":
                trailing_blank_count += 1
            else:
                break

        # snapshot formatting BEFORE deletion
        tpl = self._snapshot_template(old_block)

        # build new lines
        new_lines = pasted_text.splitlines()
        while new_lines and new_lines[-1].strip() == "":
            new_lines.pop()

        # delete old block paragraphs (so section can grow/shrink)
        for p in reversed(old_block):
            self._delete_paragraph(p)

        # insert new paragraphs before the next heading
        if not new_lines:
            self._add_line_with_template(next_heading_para, "", tpl)
        else:
            for line in new_lines:
                # keep exactly what user pasted (except removing trailing empty lines)
                self._add_line_with_template(next_heading_para, line, tpl)

        # re-add trailing blank spacer paragraphs
        for _ in range(trailing_blank_count):
            self._add_line_with_template(next_heading_para, "", tpl)

    def save(self, output_path: str):
        self.doc.save(output_path)
