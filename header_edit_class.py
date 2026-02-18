# header_edit_class.py
from docx import Document
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH


class HeaderEditor:
    HEADING_RE = re.compile(r"^[A-Z0-9 &/\-]+$")
    EMAIL_RE = re.compile(r"\b[\w\.-]+@[\w\.-]+\.\w+\b")
    PHONE_RE = re.compile(r"\+?\d[\d\-\s\(\)]{7,}\d")

    def __init__(self, resume_path: str):
        self.resume_path = resume_path
        self.doc = Document(resume_path)

    def _is_caps_header(self, text: str) -> bool:
        t = (text or "").strip()
        if not t:
            return False
        letters = [c for c in t if c.isalpha()]
        return bool(letters) and t == t.upper() and bool(self.HEADING_RE.match(t))

    def _first_caps_header(self) -> str:
        for p in self.doc.paragraphs:
            t = (p.text or "").strip()
            if self._is_caps_header(t):
                return t
        raise ValueError("Name header not found.")

    def _extract_hyperlinks(self, paragraph):
        results = []
        ns = paragraph.part.element.nsmap
        for h in paragraph._p.findall(".//w:hyperlink", ns):
            r_id = h.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            text = "".join(t.text for t in h.findall(".//w:t", ns) if t.text)
            url = None
            if r_id:
                rel = paragraph.part.rels.get(r_id)
                if rel:
                    url = rel.target_ref
            if text:
                results.append((text, url, r_id))
        return results

    def _find_contact_paragraph(self, name_header: str):
        paras = self.doc.paragraphs
        idx = None
        for i, p in enumerate(paras):
            if (p.text or "").strip().upper() == name_header.strip().upper():
                idx = i
                break
        if idx is None:
            raise ValueError("Name header paragraph not found.")

        i = idx + 1
        while i < len(paras):
            t = (paras[i].text or "").strip()
            if t and self._is_caps_header(t):
                break
            if t and (self.EMAIL_RE.search(t) or self.PHONE_RE.search(t) or "LinkedIn" in t or "GitHub" in t):
                return paras[i]
            i += 1
        raise ValueError("Contact line not found.")

    def _set_text_before_first_hyperlink(self, paragraph, new_prefix: str):
        ns = paragraph.part.element.nsmap
        children = list(paragraph._p)

        first_h_idx = None
        for i, child in enumerate(children):
            if child.tag.endswith("}hyperlink"):
                first_h_idx = i
                break

        if first_h_idx is None:
            paragraph.text = new_prefix.strip()
            return

        runs_before = [c for c in children[:first_h_idx] if c.tag.endswith("}r")]
        if not runs_before:
            paragraph.add_run(new_prefix)
            return

        first_run = runs_before[0]
        texts = first_run.findall(".//w:t", ns)
        if texts:
            texts[0].text = new_prefix
            for t in texts[1:]:
                t.text = ""
        for run in runs_before[1:]:
            for t in run.findall(".//w:t", ns):
                t.text = ""

    def _set_sep_between_two_hyperlinks(self, paragraph, sep=" • "):
        ns = paragraph.part.element.nsmap
        children = list(paragraph._p)
        hyper_idxs = [i for i, c in enumerate(children) if c.tag.endswith("}hyperlink")]
        if len(hyper_idxs) < 2:
            return

        start = hyper_idxs[0] + 1
        end = hyper_idxs[1]

        runs_between = [c for c in children[start:end] if c.tag.endswith("}r")]
        if not runs_between:
            return

        first_run = runs_between[0]
        texts = first_run.findall(".//w:t", ns)
        if texts:
            texts[0].text = sep
            for t in texts[1:]:
                t.text = ""
        for run in runs_between[1:]:
            for t in run.findall(".//w:t", ns):
                t.text = ""

    def _update_link(self, paragraph, contains_text: str, new_url: str):
        if not new_url:
            return
        ns = paragraph.part.element.nsmap
        for h in paragraph._p.findall(".//w:hyperlink", ns):
            r_id = h.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            display_text = "".join(t.text for t in h.findall(".//w:t", ns) if t.text)
            if display_text and contains_text.lower() in display_text.lower() and r_id:
                rel = paragraph.part.rels.get(r_id)
                if rel:
                    rel._target = new_url

    def get_current(self) -> dict:
        name = self._first_caps_header()
        p = self._find_contact_paragraph(name)
        txt = (p.text or "").strip()
        parts = [x.strip() for x in txt.split("•") if x.strip()]

        # best-effort parsing
        location = parts[0] if parts else ""
        phone = next((x for x in parts if self.PHONE_RE.search(x)), "")
        email = next((x for x in parts if self.EMAIL_RE.search(x)), "")

        links = self._extract_hyperlinks(p)
        linkedin_url = next((u for t, u, _ in links if "linkedin" in t.lower()), "")
        github_url = next((u for t, u, _ in links if "github" in t.lower()), "")

        return {
            "name": name,
            "location": location,
            "phone": phone,
            "email": email,
            "linkedin_url": linkedin_url,
            "github_url": github_url
        }

    def update(self, location: str, phone: str, email: str, linkedin_url: str, github_url: str):
        name = self._first_caps_header()
        p = self._find_contact_paragraph(name)

        prefix = f"{location} • {phone} • {email} • "
        self._set_text_before_first_hyperlink(p, prefix)
        self._set_sep_between_two_hyperlinks(p, " • ")
        self._update_link(p, "LinkedIn", linkedin_url)
        self._update_link(p, "GitHub", github_url)

    def save(self, output_path: str):
        self.doc.save(output_path)
