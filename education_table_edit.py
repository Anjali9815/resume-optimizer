# education_table_edit.py
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class EducationTableEditor:
    def __init__(self, resume_path: str, table_index: int = 0, row_index: int = 0):
        self.resume_path = resume_path
        self.doc = Document(resume_path)
        self.table_index = table_index
        self.row_index = row_index

    def _cell_text(self, cell) -> str:
        return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()

    def get_current(self) -> dict:
        tbl = self.doc.tables[self.table_index]
        row = tbl.rows[self.row_index]
        return {
            "left": self._cell_text(row.cells[0]),
            "right": self._cell_text(row.cells[1]),
        }

    def _copy_run_style(self, src_run, dst_run):
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size
        if src_run.font.color and src_run.font.color.rgb:
            dst_run.font.color.rgb = src_run.font.color.rgb

    def _set_cell_text_preserve_style(self, cell, new_text: str, align_right: bool = False):
        """
        Replace cell text while preserving original run formatting.
        DOES NOT change table sizing/widths.
        """
        new_text = new_text.strip()

        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        src_run = p.runs[0] if p.runs else None

        # Clear text in all runs across all paragraphs (don't delete paragraphs)
        for para in cell.paragraphs:
            for r in para.runs:
                r.text = ""

        # Write new text into first paragraph
        if p.runs:
            r0 = p.runs[0]
            r0.text = new_text
        else:
            r0 = p.add_run(new_text)

        if src_run:
            self._copy_run_style(src_run, r0)

        # Keep original paragraph formatting; only set alignment
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT

    def update(self, left_text: str, right_text: str):
        tbl = self.doc.tables[self.table_index]
        row = tbl.rows[self.row_index]

        self._set_cell_text_preserve_style(row.cells[0], left_text, align_right=False)
        self._set_cell_text_preserve_style(row.cells[1], right_text, align_right=True)

    def save(self, output_path: str):
        self.doc.save(output_path)
