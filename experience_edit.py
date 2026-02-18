# table_section_editor.py
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy


class ExperienceEditor:
    def __init__(self, resume_path: str):
        self.doc = Document(resume_path)

    # ---------- basic helpers ----------
    def _body_children(self):
        return list(self.doc._body._element)

    def _paragraph_from_elem(self, elem):
        for p in self.doc.paragraphs:
            if p._p is elem:
                return p
        return None

    def _delete_paragraph(self, paragraph):
        el = paragraph._element
        el.getparent().remove(el)
        paragraph._p = paragraph._element = None

    def _copy_run_style(self, src_run, dst_run):
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size
        if src_run.font.color and src_run.font.color.rgb:
            dst_run.font.color.rgb = src_run.font.color.rgb

    def _set_para_spacing(self, p, space_before=None, space_after=None):
        pf = p.paragraph_format
        if space_before is not None:
            pf.space_before = space_before
        if space_after is not None:
            pf.space_after = space_after

    # ---------- table location ----------
    def _find_table_pos(self, table_index: int) -> int:
        tbl_elem = self.doc.tables[table_index]._tbl
        children = self._body_children()
        for i, child in enumerate(children):
            if child is tbl_elem:
                return i
        raise ValueError("Table not found in document body.")

    def _find_next_table_elem(self, table_index: int):
        children = self._body_children()
        pos = self._find_table_pos(table_index)
        for child in children[pos + 1:]:
            if child.tag.endswith("}tbl"):
                return child
        return None

    # ---------- header read/update ----------
    def get_table_header(self, table_index: int) -> dict:
        tbl = self.doc.tables[table_index]
        row = tbl.rows[0]
        left = " ".join(p.text.strip() for p in row.cells[0].paragraphs if p.text.strip()).strip()
        right = " ".join(p.text.strip() for p in row.cells[1].paragraphs if p.text.strip()).strip()
        return {"left": left, "right": right}

    def _set_cell_text_preserve_style(self, cell, new_text: str, align_right: bool = False):
        new_text = new_text.strip()
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        src_run = p.runs[0] if p.runs else None

        for para in cell.paragraphs:
            for r in para.runs:
                r.text = ""

        if p.runs:
            r0 = p.runs[0]
            r0.text = new_text
        else:
            r0 = p.add_run(new_text)

        if src_run:
            self._copy_run_style(src_run, r0)

        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT

    def update_header(self, table_index: int, left_text: str, right_text: str):
        tbl = self.doc.tables[table_index]
        row = tbl.rows[0]
        self._set_cell_text_preserve_style(row.cells[0], left_text, align_right=False)
        self._set_cell_text_preserve_style(row.cells[1], right_text, align_right=True)

    # ---------- bullet detection ----------
    def _is_bullet_paragraph(self, p) -> bool:
        if p._p.pPr is not None and p._p.pPr.numPr is not None:
            return True
        try:
            style_name = (p.style.name or "").lower()
        except Exception:
            style_name = ""
        if "list" in style_name or "bullet" in style_name:
            return True
        return (p.text or "").lstrip().startswith("•")

    def _remove_leading_empty_paragraphs_after_table(self, table_index: int, max_remove: int = 50):
        children = self._body_children()
        pos = self._find_table_pos(table_index)

        removed = 0
        for child in children[pos + 1:]:
            if removed >= max_remove:
                break
            if child.tag.endswith("}tbl"):
                break
            if not child.tag.endswith("}p"):
                continue

            p_obj = self._paragraph_from_elem(child)
            if p_obj is None:
                continue

            if (p_obj.text or "").strip() == "":
                child.getparent().remove(child)
                removed += 1
                continue

            break

    def _remove_all_empty_paragraphs_before_elem(self, elem, max_remove: int = 50):
        if elem is None:
            return
        removed = 0
        prev = elem.getprevious()
        while prev is not None and removed < max_remove:
            if not prev.tag.endswith("}p"):
                break
            p_obj = self._paragraph_from_elem(prev)
            if p_obj is not None and (p_obj.text or "").strip() == "":
                prev.getparent().remove(prev)
                removed += 1
                prev = elem.getprevious()
                continue
            break

    def _snapshot_spacer_before_elem(self, elem):
        if elem is None:
            return None
        prev = elem.getprevious()
        while prev is not None and prev.tag.endswith("}p"):
            p_obj = self._paragraph_from_elem(prev)
            if p_obj is not None and (p_obj.text or "").strip() == "":
                style = p_obj.style
                ppr = deepcopy(p_obj._p.pPr) if (p_obj._p is not None and p_obj._p.pPr is not None) else None
                run0 = p_obj.runs[0] if p_obj.runs else None
                return {"style": style, "ppr": ppr, "run0": run0}
            break
        return None

    def _insert_one_spacer_before_elem(self, elem, spacer_tpl, fallback_run_src=None):
        if elem is None:
            return
        new_p = self.doc.add_paragraph("")
        elem.addprevious(new_p._p)

        if spacer_tpl and spacer_tpl.get("style") is not None:
            new_p.style = spacer_tpl["style"]

        if spacer_tpl and spacer_tpl.get("ppr") is not None:
            if new_p._p.pPr is not None:
                new_p._p.remove(new_p._p.pPr)
            new_p._p.insert(0, deepcopy(spacer_tpl["ppr"]))
            if new_p._p.pPr is not None and new_p._p.pPr.numPr is not None:
                new_p._p.pPr.remove(new_p._p.pPr.numPr)

        r = new_p.add_run("")
        if spacer_tpl and spacer_tpl.get("run0") is not None:
            self._copy_run_style(spacer_tpl["run0"], r)
        elif fallback_run_src is not None:
            self._copy_run_style(fallback_run_src, r)

    def get_bullets_after_table(self, table_index: int) -> list:
        children = self._body_children()
        pos = self._find_table_pos(table_index)

        bullets = []
        started = False
        for child in children[pos + 1:]:
            if child.tag.endswith("}tbl"):
                break
            if not child.tag.endswith("}p"):
                continue

            p_obj = self._paragraph_from_elem(child)
            if p_obj is None:
                continue

            txt = (p_obj.text or "").strip()
            if not txt:
                continue

            if self._is_bullet_paragraph(p_obj):
                bullets.append(p_obj)
                started = True
            else:
                if started:
                    break
        return bullets

    def list_bullet_texts(self, table_index: int) -> list[str]:
        return [(p.text or "").strip() for p in self.get_bullets_after_table(table_index)]

    def replace_all_bullets_scoped(self, table_index: int, new_bullets: list[str], next_table_override: int | None = None, keep_one_blank_line_before_next: bool = True):
        """
        Same as replace_all_bullets(), but anchors insertion to the next table INSIDE the same section list.
        - next_table_override = the next table index in the same section (experience or projects)
        """
        bullets = self.get_bullets_after_table(table_index)
        if not bullets:
            raise ValueError("No bullet block detected under this entry. Not modifying to avoid breaking layout.")

        new_bullets = [b.strip() for b in new_bullets if b.strip()]
        if not new_bullets:
            raise ValueError("No new bullets provided.")

        template_p = bullets[0]
        template_run = template_p.runs[0] if template_p.runs else None

        if template_p._p is None or template_p._p.pPr is None:
            raise ValueError("Template bullet paragraph has no pPr; cannot preserve bullet formatting.")

        template_ppr = deepcopy(template_p._p.pPr)
        template_style = template_p.style

        # ✅ Anchor: next table inside same section (if provided), else next table in doc
        if next_table_override is not None:
            next_tbl_elem = self.doc.tables[next_table_override]._tbl
        else:
            next_tbl_elem = self._find_next_table_elem(table_index)

        spacer_tpl = self._snapshot_spacer_before_elem(next_tbl_elem)

        # delete old bullets
        for p in reversed(bullets):
            self._delete_paragraph(p)

        # remove leading empty paragraphs after table to remove gap
        self._remove_leading_empty_paragraphs_after_table(table_index)

        # insert bullets in NORMAL order
        if next_tbl_elem is not None:
            for b in new_bullets:
                new_p = self.doc.add_paragraph("")
                next_tbl_elem.addprevious(new_p._p)

                if new_p._p.pPr is not None:
                    new_p._p.remove(new_p._p.pPr)
                new_p._p.insert(0, deepcopy(template_ppr))
                new_p.style = template_style
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                r0 = new_p.add_run(b)
                if template_run:
                    self._copy_run_style(template_run, r0)

            # remove space before first bullet
            new_block = self.get_bullets_after_table(table_index)
            if new_block:
                self._set_para_spacing(new_block[0], space_before=0)
                new_block[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # keep exactly one blank line before next header
            self._remove_all_empty_paragraphs_before_elem(next_tbl_elem, max_remove=50)
            if keep_one_blank_line_before_next:
                self._insert_one_spacer_before_elem(next_tbl_elem, spacer_tpl, fallback_run_src=template_run)

    def save(self, output_path: str):
        self.doc.save(output_path)
